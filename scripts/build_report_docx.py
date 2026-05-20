from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "PROJECT_REPORT.md"
OUTPUT = ROOT / "reports" / "Cloud-Native-Network-Monitoring-Proje-Raporu.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.style = "Table Grid"
    table.autofit = False
    set_cell_margins(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 12, 6),
        ("Heading 2", 13, RGBColor(46, 116, 181), 10, 5),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Cloud-Native Network Monitoring Sistemi"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(95, 108, 125)

    footer = section.footer.paragraphs[0]
    footer.add_run("Sayfa ")
    add_page_number(footer)

    return doc


def add_rich_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    tokens = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(31, 77, 120)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            paragraph.add_run(token)
    return paragraph


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        parsed.append(cells)

    if not parsed:
        return

    col_count = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(parsed):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.add_run(text)
    style_table(table)


def build_docx():
    doc = setup_document()

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Cloud-Native Network Monitoring Sistemi")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Akıllı Ağlar - SDN, Network Programlama ve Yapay Zekâ Uygulamaları")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(95, 108, 125)

    doc.add_paragraph()

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer = []
    code_buffer = []
    in_code = False

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    def flush_code():
        nonlocal code_buffer
        if code_buffer:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            for idx, code_line in enumerate(code_buffer):
                if idx:
                    paragraph.add_run("\n")
                run = paragraph.add_run(code_line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(31, 77, 120)
            code_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not stripped:
            flush_table()
            continue

        if stripped == "---":
            flush_table()
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue

        flush_table()

        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- "):
            add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped.startswith("> "):
            paragraph = add_rich_paragraph(doc, stripped[2:])
            paragraph.paragraph_format.left_indent = Inches(0.25)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(122, 90, 0)
        else:
            add_rich_paragraph(doc, stripped)

    flush_table()
    flush_code()

    OUTPUT.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
